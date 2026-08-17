# Bilingual (FR/EN) Word-valid EBIOS RM exit-report templates (python-docx) +
# docxtemplater tags. The whole document is built inside build(LANG); every
# fixed string goes through T("fr", "en"). Two files are produced:
# ebios-report-fr.docx and ebios-report-en.docx. docxtemplater tags ({...}) and
# @@IMG:...@@ markers are identical in both — the same _reportData() fills them.
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

TPL_DIR = os.path.join(os.path.dirname(__file__), "..", "app", "templates")
NAVY = RGBColor(0x1E, 0x3A, 0x5F); GREY = RGBColor(0x55, 0x60, 0x70); WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def build(LANG):
    def T(fr, en): return fr if LANG == "fr" else en

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = RGBColor(0x22, 0x2A, 0x33)
    for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11), ("Title", 26)):
        st = doc.styles[h]; st.font.color.rgb = NAVY; st.font.size = Pt(sz); st.font.bold = True
    S0 = doc.sections[0]; PW, PH = S0.page_width, S0.page_height

    def H(t, lvl=1): return doc.add_heading(t, level=lvl)
    def P(t): return doc.add_paragraph(t)
    def I(t): return doc.add_paragraph(t)
    def BULLET(t): return doc.add_paragraph(t, style="List Bullet")
    def CENTER(p): p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return p

    def _grey(run): run.font.size = Pt(8.5); run.font.color.rgb = GREY; run.bold = True
    def _seq_caption(prefix, seqname, title):
        p = CENTER(doc.add_paragraph())
        _grey(p.add_run(prefix + " "))
        r = p.add_run(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r._r.append(fc); _grey(r)
        r = p.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' SEQ %s \\* ARABIC ' % seqname; r._r.append(it); _grey(r)
        r = p.add_run(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate'); r._r.append(fc); _grey(r)
        _grey(p.add_run("1"))
        r = p.add_run(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end'); r._r.append(fc); _grey(r)
        if title: _grey(p.add_run(" — " + title))
        return p
    def CAPT(title): return _seq_caption(T("Tableau", "Table"), "tbl", title)
    def CAPF(title): return _seq_caption(T("Figure", "Figure"), "fig", title)

    def _field(p, instr, default="", style_run=None):
        for kind, val in [('begin', None), ('instr', instr), ('separate', None), ('text', default), ('end', None)]:
            r = p.add_run()
            if kind == 'instr':
                e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = val; r._r.append(e)
            elif kind == 'text':
                r.text = val
            else:
                e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), kind)
                if kind == 'begin': e.set(qn('w:dirty'), 'true')
                r._r.append(e)
            if style_run: style_run(r)
        return p

    def toc():
        _field(doc.add_paragraph(), ' TOC \\o "1-3" \\h \\z \\u ',
               T("La table des matières sera générée à l'ouverture (clic droit → Mettre à jour les champs).",
                 "The table of contents will be generated on opening (right-click → Update field)."))

    def _set_footer(sec, title):
        sec.footer.is_linked_to_previous = False
        p = sec.footer.paragraphs[0]; p.text = ""
        usable = sec.page_width - sec.left_margin - sec.right_margin
        right_pos = sec.page_width - sec.left_margin - 109728
        p.paragraph_format.tab_stops.add_tab_stop(int(usable / 2), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(int(right_pos), WD_TAB_ALIGNMENT.RIGHT)
        def small(r): r.font.size = Pt(8.5); r.font.color.rgb = GREY
        small(p.add_run("\t" + title)); small(p.add_run("\t"))
        _field(p, ' PAGE ', "1", style_run=small)

    _cur = ["portrait"]
    def _switch(landscape):
        want = "landscape" if landscape else "portrait"
        if _cur[0] == want: return
        s = doc.add_section(WD_SECTION.NEW_PAGE)
        if landscape: s.orientation = WD_ORIENT.LANDSCAPE; s.page_width, s.page_height = PH, PW
        else: s.orientation = WD_ORIENT.PORTRAIT; s.page_width, s.page_height = PW, PH
        _cur[0] = want
    def land(): _switch(True)
    def portrait(): _switch(False)

    _bmk = [0]
    def bookmark(p, name):
        _bmk[0] += 1; bid = str(_bmk[0])
        a = OxmlElement('w:bookmarkStart'); a.set(qn('w:id'), bid); a.set(qn('w:name'), name)
        b = OxmlElement('w:bookmarkEnd'); b.set(qn('w:id'), bid); p._p.insert(0, a); p._p.append(b)
    def link(p, anchor, text):
        h = OxmlElement('w:hyperlink'); h.set(qn('w:anchor'), anchor)
        r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
        c = OxmlElement('w:color'); c.set(qn('w:val'), '1E3A5F'); u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
        rpr.append(c); rpr.append(u); r.append(rpr)
        t = OxmlElement('w:t'); t.text = text; r.append(t); h.append(r); p._p.append(h)
    def no_number(p):
        p._p.get_or_add_pPr().append(parse_xml('<w:numPr %s><w:ilvl w:val="0"/><w:numId w:val="0"/></w:numPr>' % nsdecls('w')))

    def cartouche():
        tb = doc.add_table(rows=0, cols=2); tblPr = tb._tbl.tblPr
        tblPr.append(parse_xml('<w:tblpPr %s w:vertAnchor="margin" w:horzAnchor="margin" w:tblpXSpec="center" w:tblpYSpec="bottom"/>' % nsdecls('w')))
        tblPr.append(parse_xml('<w:tblW %s w:w="3800" w:type="pct"/>' % nsdecls('w')))
        tblPr.append(parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '</w:tblBorders>' % nsdecls('w')))
        def row(label, value, value_bold=False):
            cells = tb.add_row().cells; lc, vc = cells[0], cells[1]
            lc._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:val="clear" w:fill="EEF2F7"/>' % nsdecls('w')))
            lc._tc.get_or_add_tcPr().append(parse_xml('<w:tcW %s w:w="1700" w:type="dxa"/>' % nsdecls('w')))
            r = lc.paragraphs[0].add_run(label); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY
            r = vc.paragraphs[0].add_run(value); r.font.size = Pt(9.5); r.bold = value_bold
            return cells
        row(T("Rédacteur", "Author"), "{cart_redacteur}")
        row(T("Contributeurs", "Contributors"), "{cart_contributeurs}")
        row(T("Version", "Version"), "{cart_version}")
        row(T("Date", "Date"), "{cart_date}")
        row(T("Classification", "Classification"), "{cart_classification}", value_bold=True)
        hc = tb.add_row().cells; merged = hc[0].merge(hc[1])
        merged._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:val="clear" w:fill="1E3A5F"/>' % nsdecls('w')))
        r = merged.paragraphs[0].add_run(T("Validation", "Approval")); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
        row(T("Date", "Date"), "")
        sig = row(T("Nom / signature", "Name / signature"), "")
        sig[1].add_paragraph(""); sig[1].add_paragraph("")
        return tb

    def IMG(marker, caption):
        CENTER(P(marker)); CAPF(caption)

    def autofit(tb):
        tb.autofit = True; tb.allow_autofit = True; tblPr = tb._tbl.tblPr
        lay = tblPr.find(qn('w:tblLayout'))
        if lay is None: lay = OxmlElement('w:tblLayout'); tblPr.append(lay)
        lay.set(qn('w:type'), 'autofit')
        w = tblPr.find(qn('w:tblW'))
        if w is None: w = OxmlElement('w:tblW'); tblPr.append(w)
        w.set(qn('w:w'), '0'); w.set(qn('w:type'), 'auto')
        for row in tb.rows:
            for cell in row.cells:
                tcW = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
                if tcW is None: tcW = OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
                tcW.set(qn('w:w'), '0'); tcW.set(qn('w:type'), 'auto')
        return tb

    def _table(headers, loop, fields, caption=None):
        tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
        for i, hh in enumerate(headers): tb.rows[0].cells[i].paragraphs[0].add_run(hh).bold = True
        cells = tb.add_row().cells
        for i, f in enumerate(fields):
            tag = "{" + f + "}"
            if i == 0: tag = "{#" + loop + "}" + tag
            if i == len(fields) - 1: tag = tag + "{/" + loop + "}"
            cells[i].paragraphs[0].add_run(tag)
        autofit(tb)
        if caption: CAPT(caption)
        return tb

    def note_box(title, body):
        tb = doc.add_table(rows=1, cols=2); tblPr = tb._tbl.tblPr
        tblPr.append(parse_xml('<w:tblW %s w:w="5000" w:type="pct"/>' % nsdecls('w')))
        tblPr.append(parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:left w:val="single" w:sz="18" w:space="0" w:color="1E3A5F"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="C7D2DD"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>' % nsdecls('w')))
        icon_cell, txt_cell = tb.cell(0, 0), tb.cell(0, 1)
        for c in (icon_cell, txt_cell):
            c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:val="clear" w:fill="F4F7FA"/>' % nsdecls('w')))
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        icon_cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcW %s w:w="540" w:type="dxa"/>' % nsdecls('w')))
        ip = icon_cell.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ir = ip.add_run("💡"); ir.font.size = Pt(16)
        p = txt_cell.paragraphs[0]
        r = p.add_run(title[:1].upper() + title[1:] + " : "); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(9.5)
        r = p.add_run(body); r.font.size = Pt(9.5)
        return p

    def SUB(htext, lvl, itext, headers, loop, fields, wide, caption=None):
        land() if wide else portrait()
        H(htext, lvl)
        if itext: I(itext)
        _table(headers, loop, fields, caption or htext)

    # ===== COVER =====
    for _ in range(3): P("")
    CENTER(doc.add_paragraph(style=doc.styles["Title"])).add_run(T("Rapport d'analyse des risques", "Risk Analysis Report"))
    sub = CENTER(doc.add_paragraph()); r = sub.add_run(T("Méthode EBIOS Risk Manager (ANSSI)", "EBIOS Risk Manager method (ANSSI)"))
    r.font.size = Pt(13); r.font.color.rgb = GREY
    for _ in range(3): P("")
    org = CENTER(doc.add_paragraph()); r = org.add_run("{contexte_societe}"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    peri = CENTER(doc.add_paragraph()); r = peri.add_run(T("Périmètre de l'étude : {contexte_objet}", "Scope of the study: {contexte_objet}")); r.font.size = Pt(12)
    cartouche()
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    no_number(H(T("Table des matières", "Table of contents")))
    toc()
    doc.add_page_break()

    # ===== EXECUTIVE SUMMARY =====
    no_number(H(T("Synthèse managériale", "Executive summary")))
    P(T("La présente analyse de risques porte sur {contexte_objet} de l'organisation {contexte_societe}. Elle applique "
        "la méthode EBIOS Risk Manager de l'ANSSI, structurée en cinq ateliers :",
        "This risk analysis covers {contexte_objet} of {contexte_societe}. It applies ANSSI's EBIOS Risk Manager method, "
        "structured in five workshops:"))
    BULLET(T("Atelier 1 — cadrage et socle de sécurité", "Workshop 1 — scope and security baseline"))
    BULLET(T("Atelier 2 — identification des sources de risque et des objectifs visés", "Workshop 2 — identification of risk origins and target objectives"))
    BULLET(T("Atelier 3 — construction des scénarios stratégiques", "Workshop 3 — building strategic scenarios"))
    BULLET(T("Atelier 4 — élaboration des scénarios opérationnels", "Workshop 4 — developing operational scenarios"))
    BULLET(T("Atelier 5 — traitement et acceptation du risque", "Workshop 5 — risk treatment and acceptance"))
    P(T("L'objectif est d'éclairer les décisions de la direction sur les risques à traiter en priorité.",
        "Its purpose is to inform management decisions on the risks to address as a priority."))
    P(T("À l'issue de l'étude, la cartographie des risques résiduels fait apparaître :",
        "At the end of the study, the residual risk map shows:"))
    BULLET(T("{dist_eleve} risque(s) de niveau élevé", "{dist_eleve} high-level risk(s)"))
    BULLET(T("{dist_moyen} risque(s) de niveau moyen", "{dist_moyen} medium-level risk(s)"))
    BULLET(T("{dist_faible} risque(s) de niveau faible", "{dist_faible} low-level risk(s)"))
    P(T("Le niveau de couverture du socle de sécurité est évalué à {socle_avg} %. {mes_todo} mesure(s) restent à "
        "mettre en œuvre dans le cadre du PACS (Plan d'Amélioration Continue de la Sécurité).",
        "The security baseline coverage is assessed at {socle_avg} %. {mes_todo} measure(s) remain to be implemented "
        "under the PACS (Security Continuous Improvement Plan)."))
    P(T("Appréciation générale et priorités retenues par la direction : ……………………………………………………………………………………",
        "Management's overall assessment and selected priorities: ……………………………………………………………………………………"))
    P(T("Décision d'acceptation des risques résiduels : ……………………………………………………………………………………",
        "Residual risk acceptance decision: ……………………………………………………………………………………"))
    doc.add_page_break()

    # ===== WORKSHOP 1 =====
    portrait()
    H(T("Cadrage et socle de sécurité", "Scope and security baseline"))
    I(T("L'objectif de cette phase de l'analyse est de délimiter le périmètre métier et technique de l'étude : "
        "recenser les valeurs métier (missions, processus et informations essentiels), les rattacher à leurs biens "
        "supports (composants du SI qui les portent), caractériser les événements redoutés et apprécier leur gravité, "
        "puis établir le socle de sécurité.",
        "The purpose of this phase of the analysis is to define the business and technical scope of the study: list "
        "the business values (essential missions, processes and information), link them to their supporting assets "
        "(the IS components that carry them), characterise the feared events and assess their severity, then establish "
        "the security baseline."))
    I(T("L'établissement du socle de sécurité consiste à mesurer les écarts du système d'information par rapport au "
        "référentiel de sécurité retenu — au choix les 42 mesures d'hygiène de l'ANSSI ou la norme ISO/IEC 27001. Ce "
        "socle constitue le niveau de sécurité de base attendu : l'analyse de risque EBIOS Risk Manager a vocation à "
        "être conduite dans un contexte où il est déjà suffisamment couvert, l'étude se concentrant alors sur les "
        "risques résiduels qui subsistent au-delà de ce socle.",
        "Establishing the security baseline means measuring the gaps of the information system against the chosen "
        "security framework — either ANSSI's 42 cyber-hygiene measures or the ISO/IEC 27001 standard. This baseline is "
        "the expected minimum security level: the EBIOS Risk Manager analysis is meant to be conducted in a context "
        "where it is already sufficiently covered, the study then focusing on the residual risks that remain beyond "
        "this baseline."))
    H(T("Contexte et objectifs de l'étude", "Context and objectives of the study"), 2)
    P(T("Objet de l'étude : {contexte_objet}", "Subject of the study: {contexte_objet}"))
    P(T("Réglementation et référentiels applicables : {contexte_reglementation}", "Applicable regulations and frameworks: {contexte_reglementation}"))
    P(T("Socle de sécurité retenu : {contexte_socle}", "Selected security baseline: {contexte_socle}"))
    P(T("Éléments de contexte complémentaires (enjeux, contraintes, hypothèses) : ……………………………………………………………",
        "Additional context (issues, constraints, assumptions): ……………………………………………………………"))
    note_box(T("valeur métier et bien support", "business value and supporting asset"),
        T("une valeur métier est ce que l'organisation doit protéger — une mission, un processus ou une information "
          "essentiels ; un bien support est l'élément concret du système d'information qui la porte (serveur, application, "
          "réseau, local, intervenant…). Une même valeur métier peut reposer sur plusieurs biens supports, et c'est "
          "généralement le bien support qui constitue la cible technique d'une attaque.",
          "a business value is what the organisation must protect — an essential mission, process or piece of "
          "information; a supporting asset is the concrete IS component that carries it (server, application, network, "
          "premises, staff…). A single business value may rely on several supporting assets, and it is generally the "
          "supporting asset that is the technical target of an attack."))
    SUB(T("Valeurs métier", "Business values"), 2,
        T("Une valeur métier est une information ou un processus essentiel à l'organisation, qu'il convient de protéger ; elle est rattachée à un responsable métier.",
          "A business value is essential information or a process for the organisation, which must be protected; it is assigned to a business owner."),
        [T("Abrév.", "Abbr."), T("Nom", "Name"), T("Nature", "Nature"), T("Description", "Description"), T("Responsable", "Owner")],
        "vm", ["id", "nom", "nature", "description", "responsable"], True,
        T("Valeurs métier de l'étude", "Business values of the study"))
    SUB(T("Biens supports", "Supporting assets"), 2,
        T("Un bien support est un composant du SI (serveur, application, réseau, local, intervenant…) sur lequel repose une ou plusieurs valeurs métier.",
          "A supporting asset is an IS component (server, application, network, premises, staff…) on which one or more business values rely."),
        [T("Abrév.", "Abbr."), T("Nom", "Name"), T("Type", "Type"), T("Valeur métier", "Business value"), T("Localisation", "Location"), T("Propriétaire", "Owner")],
        "bs", ["id", "nom", "type", "vm", "localisation", "proprietaire"], False,
        T("Biens supports rattachés aux valeurs métier", "Supporting assets linked to business values"))
    land()
    H(T("Événements redoutés", "Feared events"), 2)
    I(T("Un événement redouté est une atteinte à une valeur métier (disponibilité, intégrité, confidentialité, "
        "traçabilité) et les impacts qui en découlent (missions, humains, financiers, juridiques, image). Le niveau de "
        "gravité retenu pour chaque événement redouté correspond au maximum des niveaux atteints sur les différents "
        "critères de l'échelle de gravité (financier, réputation, réglementaire, données personnelles, opérationnel).",
        "A feared event is an attack on a business value (availability, integrity, confidentiality, traceability) and "
        "the resulting impacts (missions, people, financial, legal, image). The severity retained for each feared event "
        "is the maximum of the levels reached across the various criteria of the severity scale (financial, reputation, "
        "regulatory, personal data, operational)."))
    pg = P(T("L'échelle de gravité utilisée est détaillée en ", "The severity scale used is detailed in "))
    link(pg, "annexe_gravite", T("Annexe A — Échelle de gravité", "Appendix A — Severity scale")); pg.add_run(".")
    _table([T("Abrév.", "Abbr."), T("Valeur métier", "Business value"), T("Événement redouté", "Feared event"), T("Impacts", "Impacts"), T("Gravité", "Severity")],
           "er", ["id", "vm", "evenement", "impacts", "gravite"],
           T("Événements redoutés et niveaux de gravité", "Feared events and severity levels"))
    H(T("Socle de sécurité", "Security baseline"), 2)
    I(T("Le socle de sécurité regroupe les exigences des référentiels applicables et leur état d'application. Le taux de "
        "couverture global est la moyenne des taux de conformité (en %) des exigences évaluées.",
        "The security baseline groups the requirements of the applicable frameworks and their state of application. "
        "The overall coverage rate is the average of the compliance rates (in %) of the assessed requirements."))
    P(T("Niveau de couverture global du socle évalué dans cette étude : {socle_avg} %.",
        "Overall baseline coverage assessed in this study: {socle_avg} %."))
    _table([T("Réf. exigence", "Req. ref."), T("Thème", "Theme"), T("Mesure", "Measure"), T("Conformité", "Compliance"), T("Écart", "Gap")],
           "socle", ["ref", "theme", "mesure", "conformite", "ecart"],
           T("Évaluation du socle de sécurité", "Security baseline assessment"))
    H(T("Mesures prévues pour le socle", "Planned baseline measures"), 3)
    I(T("Pour les exigences partiellement ou non couvertes, les mesures correctives déjà prévues alimentent le PACS "
        "(Plan d'Amélioration Continue de la Sécurité, chapitre Traitement du risque).",
        "For partially or non-covered requirements, the corrective measures already planned feed the PACS "
        "(Security Continuous Improvement Plan, Risk treatment chapter)."))
    _table([T("Réf. exigence", "Req. ref."), T("Mesure / exigence", "Measure / requirement"), T("Mesures prévues", "Planned measures")],
           "socle_planned", ["ref", "mesure", "mesures_prevues"],
           T("Mesures prévues pour combler les écarts du socle", "Measures planned to close the baseline gaps"))

    # ===== WORKSHOP 2 =====
    portrait()
    H(T("Sources de risque et objectifs visés", "Risk origins and target objectives"))
    I(T("L'objectif de cette phase de l'analyse est d'identifier « qui ou quoi » pourrait porter atteinte aux valeurs "
        "métier et « pourquoi », puis de retenir les couples source de risque / objectif visé les plus pertinents.",
        "The purpose of this phase of the analysis is to identify “who or what” could harm the business values "
        "and “why”, then retain the most relevant risk-origin / target-objective pairs."))
    SUB(T("Sources de risque", "Risk origins"), 2,
        T("Une source de risque (SR) est un acteur — humain ou non — susceptible de générer un risque : cybercriminel, "
          "concurrent, État, hacktiviste, employé malveillant ou négligent, prestataire…",
          "A risk origin (RO) is an actor — human or not — likely to generate a risk: cybercriminal, competitor, state, "
          "hacktivist, malicious or negligent employee, service provider…"),
        [T("Abrév.", "Abbr."), T("Source de risque", "Risk origin")], "sr_list", ["id", "nom"], False,
        T("Sources de risque identifiées", "Identified risk origins"))
    SUB(T("Objectifs visés", "Target objectives"), 2,
        T("Un objectif visé (OV) est la finalité poursuivie par la source de risque : espionnage, sabotage, gain financier, "
          "lucratif, déstabilisation, défi…",
          "A target objective (TO) is the aim pursued by the risk origin: espionage, sabotage, financial gain, "
          "destabilisation, challenge…"),
        [T("Abrév.", "Abbr."), T("Objectif visé", "Target objective")], "ov_list", ["id", "nom"], False,
        T("Objectifs visés identifiés", "Identified target objectives"))
    SUB(T("Couples Source de risque / Objectif visé évalués", "Risk origin / Target objective pairs assessed"), 2,
        T("La pertinence d'un couple SR/OV est la somme de trois critères notés de 0 à 4 — motivation, ressources et "
          "activité de la source (score de 0 à 12). La priorité en découle : P1 si le score est strictement supérieur "
          "à 7, P2 s'il est compris entre 5 et 7, « Non retenu » entre 3 et 4, « Écarté » entre 1 et 2. Le tableau "
          "ci-dessous présente l'ensemble des couples évalués avec leur priorité : seuls les couples P1 et P2 sont "
          "retenus pour la suite de l'étude, les couples « Non retenu » et « Écarté » étant exclus.",
          "The relevance of an RO/TO pair is the sum of three criteria rated 0 to 4 — motivation, resources and "
          "activity of the origin (score 0 to 12). The priority follows: P1 if the score is strictly above 7, P2 if "
          "between 5 and 7, “Not retained” between 3 and 4, “Discarded” between 1 and 2. The table "
          "below lists all assessed pairs with their priority: only P1 and P2 pairs are retained for the rest of the "
          "study, “Not retained” and “Discarded” pairs being excluded."),
        [T("Couple", "Pair"), T("Source", "Origin"), T("Objectif visé", "Target objective"), T("Motiv.", "Motiv."), T("Ress.", "Res."), T("Activ.", "Act."), T("Pertinence", "Relevance"), T("Priorité", "Priority"), T("Justification", "Justification")],
        "srov", ["couple", "sr", "ov", "motivation", "ressources", "activite", "pertinence", "priorite", "justification"], True,
        T("Couples SR/OV évalués et priorisés", "RO/TO pairs assessed and prioritised"))

    # ===== WORKSHOP 3 =====
    portrait()
    H(T("Scénarios stratégiques", "Strategic scenarios"))
    I(T("L'objectif de cette phase de l'analyse est, à partir des couples SR/OV retenus, d'établir la cartographie de la "
        "menace numérique de l'écosystème (parties prenantes : prestataires, partenaires, fournisseurs…), puis de "
        "construire les scénarios stratégiques — chemins d'attaque de haut niveau, souvent via l'écosystème, vers les "
        "valeurs métier. Chaque partie prenante est évaluée selon quatre critères : dépendance, pénétration, maturité "
        "cyber et confiance.",
        "The purpose of this phase of the analysis is, from the retained RO/TO pairs, to map the digital threat of the "
        "ecosystem (stakeholders: service providers, partners, suppliers…), then build the strategic scenarios — "
        "high-level attack paths, often via the ecosystem, towards the business values. Each stakeholder is assessed "
        "against four criteria: dependency, penetration, cyber maturity and trust."))
    SUB(T("Parties prenantes de l'écosystème", "Ecosystem stakeholders"), 2, None,
        [T("Abrév.", "Abbr."), T("Nom", "Name"), T("Catégorie", "Category")], "pp", ["id", "nom", "categorie"], False,
        T("Parties prenantes de l'écosystème", "Ecosystem stakeholders"))
    SUB(T("Évaluation des parties prenantes (menace)", "Stakeholder assessment (threat)"), 2,
        T("Le niveau de menace combine l'exposition (dépendance × pénétration) et la fiabilité (maturité × confiance) de "
          "chaque partie prenante : menace = (dépendance × pénétration) / (maturité × confiance). Le niveau d'exposition "
          "qui en résulte (faible à critique) est indiqué pour chaque partie prenante.",
          "The threat level combines exposure (dependency × penetration) and reliability (maturity × trust) of each "
          "stakeholder: threat = (dependency × penetration) / (maturity × trust). The resulting exposure level (low to "
          "critical) is given for each stakeholder."),
        [T("Nom", "Name"), T("Dépend.", "Depend."), T("Pénétr.", "Penetr."), T("Matur.", "Matur."), T("Confiance", "Trust"), T("Menace", "Threat"), T("Exposition", "Exposure")],
        "pp_eval", ["nom", "dependance", "penetration", "maturite", "confiance", "menace", "exposition"], False,
        T("Évaluation de la menace portée par les parties prenantes", "Assessment of the threat posed by stakeholders"))
    land()
    H(T("Cartographie de la menace de l'écosystème", "Ecosystem threat map"), 2)
    I(T("Représentation de la menace portée par les parties prenantes, avant puis après application des mesures sur l'écosystème.",
        "Representation of the threat posed by stakeholders, before and after applying ecosystem measures."))
    note_box(T("comment lire la cartographie", "how to read the map"),
        T("chaque partie prenante est un cercle positionné par rapport au centre : plus elle est proche du centre, plus "
          "la menace qu'elle représente est élevée. La couleur traduit sa fiabilité (maturité × confiance), du rouge "
          "(faible) au vert (élevée) ; la taille du cercle reflète son exposition (dépendance × pénétration). Les zones "
          "concentriques distinguent la zone de veille (extérieure), de contrôle, puis de danger (centrale).",
          "each stakeholder is a circle positioned relative to the centre: the closer to the centre, the higher the "
          "threat it represents. The colour reflects its reliability (maturity × trust), from red (low) to green (high); "
          "the size of the circle reflects its exposure (dependency × penetration). The concentric zones distinguish "
          "the watch zone (outer), control, then danger (central)."))
    IMG("@@IMG:pp_map_initial@@", T("Cartographie de la menace de l'écosystème — avant traitement", "Ecosystem threat map — before treatment"))
    IMG("@@IMG:pp_map_residual@@", T("Cartographie de la menace de l'écosystème — après traitement", "Ecosystem threat map — after treatment"))
    SUB(T("Scénarios stratégiques identifiés", "Identified strategic scenarios"), 2,
        T("{ss_count} scénario(s) stratégique(s) ont été identifiés à l'issue de cet atelier ; ils sont décrits dans le tableau suivant.",
          "{ss_count} strategic scenario(s) were identified at the end of this workshop; they are described in the table below."),
        [T("Réf.", "Ref."), T("Scénario stratégique", "Strategic scenario"), T("Partie prenante", "Stakeholder"), T("Bien support", "Supporting asset"), T("Événement redouté", "Feared event")],
        "ss", ["id", "scenario", "pp", "bs", "er"], True,
        T("Scénarios stratégiques identifiés", "Identified strategic scenarios"))

    # ===== WORKSHOP 4 =====
    portrait()
    H(T("Scénarios opérationnels", "Operational scenarios"))
    I(T("L'objectif de cette phase de l'analyse est de détailler techniquement chaque scénario stratégique en une "
        "séquence d'actions élémentaires (chemin d'attaque). Dans le contexte de cette étude, les actions ne sont pas regroupées selon les "
        "quatre phases génériques d'EBIOS RM (connaître, rentrer, trouver, exploiter) mais selon les phases de la kill "
        "chain du référentiel MITRE ATT&CK (reconnaissance, accès initial, exécution, persistance, élévation de "
        "privilèges, mouvement latéral, exfiltration / impact), plus représentatives des modes opératoires observés.",
        "The purpose of this phase of the analysis is to technically detail each strategic scenario as a sequence of "
        "elementary actions (attack path). In the context of this study, the actions are not grouped according to the "
        "four generic EBIOS RM phases (know, enter, find, exploit) but according to the kill-chain phases of the MITRE "
        "ATT&CK framework (reconnaissance, initial access, execution, persistence, privilege escalation, lateral "
        "movement, exfiltration / impact), which are more representative of the observed modes of operation."))
    note_box(T("taux de faiblesse", "weakness rate"),
        T("pour chaque scénario opérationnel, on évalue un taux de faiblesse. Pour chaque phase du chemin d'attaque, on "
          "apprécie l'efficacité des mesures déjà en place (Absent, Partiel, Efficace), qui réduit la probabilité que la "
          "phase aboutisse. Le taux de faiblesse du scénario agrège ces efficacités selon la formule : "
          "taux = max(0 ; 2 × n(Absent) + n(Partiel) − 2 × n(Efficace)) / (2 × n(phases)). Un taux proche de 1 traduit un "
          "chemin d'attaque peu entravé par les mesures existantes.",
          "for each operational scenario, a weakness rate is assessed. For each phase of the attack path, the "
          "effectiveness of the measures already in place is appreciated (Absent, Partial, Effective), which reduces "
          "the probability that the phase succeeds. The scenario's weakness rate aggregates these effectiveness levels "
          "using the formula: rate = max(0 ; 2 × n(Absent) + n(Partial) − 2 × n(Effective)) / (2 × n(phases)). A rate "
          "close to 1 reflects an attack path little hindered by existing measures."))
    note_box(T("vraisemblance", "likelihood"),
        T("le taux de faiblesse est converti en vraisemblance opérationnelle (V op) : V op = 4 si le taux ≥ 0,7 ; 3 si "
          "≥ 0,4 ; 2 si ≥ 0,2 ; 1 sinon. La vraisemblance porte sur le scénario stratégique : elle est égale au maximum "
          "des vraisemblances opérationnelles des scénarios opérationnels qui le composent.",
          "the weakness rate is converted into an operational likelihood (Lo): Lo = 4 if the rate ≥ 0.7; 3 if ≥ 0.4; "
          "2 if ≥ 0.2; 1 otherwise. The likelihood applies to the strategic scenario: it equals the maximum of the "
          "operational likelihoods of the operational scenarios that compose it."))
    land()
    doc.add_paragraph("{#ss_groups}")
    H("{ss_label}", 2)
    P("{ss_intro_pre}")
    doc.add_paragraph("{#er_multi}")
    P(T("Il pourrait aboutir aux événements redoutés suivants :", "It could lead to the following feared events:"))
    BULLET("{#er_list}{er}{/er_list}")
    doc.add_paragraph("{/er_multi}")
    doc.add_paragraph("{#er_single}")
    P(T("Il pourrait aboutir à l'événement redouté suivant : {er_one}.", "It could lead to the following feared event: {er_one}."))
    doc.add_paragraph("{/er_single}")
    P("{ss_intro_post}")
    doc.add_paragraph("{#sops}")
    H(T("Scénario opérationnel {sop_label}", "Operational scenario {sop_label}"), 3)
    P("{taux_phrase}")
    _tbl = doc.add_table(rows=1, cols=4); _tbl.style = "Light Grid Accent 1"
    for i, hh in enumerate([T("Phase", "Phase"), T("Action élémentaire", "Elementary action"), T("Contrôle / mesure", "Control / measure"), T("Efficacité", "Effectiveness")]):
        _tbl.rows[0].cells[i].paragraphs[0].add_run(hh).bold = True
    _c = _tbl.add_row().cells
    for i, tag in enumerate(["{#steps}{phase}", "{action}", "{controle}", "{efficacite}{/steps}"]):
        _c[i].paragraphs[0].add_run(tag)
    autofit(_tbl)
    CAPT(T("Chemin d'attaque du scénario opérationnel {sop_label}", "Attack path of operational scenario {sop_label}"))
    doc.add_paragraph("{/sops}")
    doc.add_paragraph("{/ss_groups}")

    # ===== WORKSHOP 5 =====
    portrait()
    H(T("Traitement du risque", "Risk treatment"))
    I(T("L'objectif de cette phase de l'analyse est de synthétiser les risques (gravité × vraisemblance), de définir le "
        "plan d'amélioration continue de la sécurité (PACS) puis de réapprécier les risques résiduels après mise en "
        "œuvre des mesures, en vue de leur acceptation par la direction.",
        "The purpose of this phase of the analysis is to synthesise the risks (severity × likelihood), define the "
        "security continuous improvement plan (PACS) then re-assess the residual risks after the measures are "
        "implemented, with a view to their acceptance by management."))
    pr = note_box(T("niveau de risque", "risk level"),
        T("le niveau de chaque risque résulte du croisement de la gravité de l'événement redouté et de la "
          "vraisemblance du scénario, selon la matrice d'acceptabilité détaillée en ",
          "the level of each risk results from crossing the severity of the feared event and the likelihood of the "
          "scenario, according to the acceptability matrix detailed in "))
    link(pr, "annexe_risque", T("Annexe B — Niveaux de risque", "Appendix B — Risk levels"))
    r = pr.add_run("."); r.font.size = Pt(9.5)
    H(T("Cartographie des risques", "Risk map"), 2)
    I(T("Positionnement des risques (gravité × vraisemblance) avant et après traitement.",
        "Positioning of risks (severity × likelihood) before and after treatment."))
    note_box(T("comment lire la matrice", "how to read the matrix"),
        T("chaque risque est positionné selon la gravité de son événement redouté (axe vertical) et la vraisemblance de "
          "son scénario (axe horizontal). La couleur de la case indique le niveau de risque résultant — faible, moyen ou "
          "élevé — selon la matrice d'acceptabilité retenue (Annexe B). Plus un risque se situe vers le haut et la droite, "
          "plus il est critique et prioritaire à traiter.",
          "each risk is positioned according to the severity of its feared event (vertical axis) and the likelihood of "
          "its scenario (horizontal axis). The colour of the cell indicates the resulting risk level — low, medium or "
          "high — according to the acceptability matrix retained (Appendix B). The further a risk is towards the top "
          "right, the more critical and the higher priority to treat."))
    IMG("@@IMG:risk_map_initial@@", T("Cartographie des risques — avant traitement", "Risk map — before treatment"))
    IMG("@@IMG:risk_map_residual@@", T("Cartographie des risques — après traitement", "Risk map — after treatment"))
    portrait()
    H(T("Risques initiaux", "Initial risks"), 2)
    P(T("La cartographie initiale des risques (avant traitement) se répartit comme suit :",
        "The initial risk map (before treatment) breaks down as follows:"))
    BULLET(T("{dist_init_eleve} risque(s) de niveau élevé", "{dist_init_eleve} high-level risk(s)"))
    BULLET(T("{dist_init_moyen} risque(s) de niveau moyen", "{dist_init_moyen} medium-level risk(s)"))
    BULLET(T("{dist_init_faible} risque(s) de niveau faible", "{dist_init_faible} low-level risk(s)"))
    land()
    _table([T("Réf.", "Ref."), T("Scénario", "Scenario"), T("Gravité", "Severity"), T("Vraisemblance", "Likelihood"), T("Niveau initial", "Initial level")],
           "risks_init", ["id", "scenario", "gravite", "vInit", "riskInit"],
           T("Risques initiaux (avant traitement)", "Initial risks (before treatment)"))
    SUB(T("Plan de traitement (PACS)", "Treatment plan (PACS)"), 2,
        T("Le PACS regroupe les actions proposées pour réduire les risques.",
          "The PACS groups the actions proposed to reduce the risks."),
        [T("N°", "No."), T("Mesure", "Measure"), T("Origine", "Origin"), T("Coût", "Cost"), T("Responsable", "Owner"), T("Échéance", "Due date"), T("Statut", "Status")],
        "measures", ["id", "mesure", "origine", "cout", "responsable", "echeance", "statut"], True,
        T("Plan d'amélioration continue de la sécurité (actions en cours et à venir)", "Security continuous improvement plan (current and upcoming actions)"))
    portrait()
    H(T("Risques résiduels", "Residual risks"), 2)
    P(T("Après mise en œuvre du plan de traitement, les risques résiduels se répartissent comme suit :",
        "After implementing the treatment plan, the residual risks break down as follows:"))
    BULLET(T("{dist_eleve} risque(s) de niveau élevé", "{dist_eleve} high-level risk(s)"))
    BULLET(T("{dist_moyen} risque(s) de niveau moyen", "{dist_moyen} medium-level risk(s)"))
    BULLET(T("{dist_faible} risque(s) de niveau faible", "{dist_faible} low-level risk(s)"))
    _table([T("Réf.", "Ref."), T("Scénario", "Scenario"), T("Évolution (initial → résiduel)", "Change (initial → residual)")],
           "risks_resid", ["id", "scenario", "reduction"],
           T("Risques résiduels et effet du traitement", "Residual risks and effect of treatment"))
    note_box(T("acceptation du risque résiduel", "residual risk acceptance"),
        T("le risque résiduel est le niveau de risque qui subsiste une fois les mesures de traitement mises en œuvre ; il "
          "ne peut jamais être totalement nul. La direction doit statuer explicitement sur son acceptation, arbitrer les "
          "risques jugés encore trop élevés et tracer cette décision, qui engage sa responsabilité.",
          "the residual risk is the level of risk that remains once the treatment measures are implemented; it can never "
          "be entirely nil. Management must explicitly decide on its acceptance, arbitrate the risks deemed still too "
          "high, and record this decision, which engages its responsibility."))

    # ===== APPENDICES =====
    land()
    no_number(H(T("Annexes", "Appendices")))
    ha = doc.add_heading(T("Annexe A — Échelle de gravité", "Appendix A — Severity scale"), level=2); no_number(ha); bookmark(ha, "annexe_gravite")
    I(T("Échelle de gravité retenue pour cette étude, adaptée au contexte de l'organisation ; chaque événement redouté y est positionné (cf. Événements redoutés).",
        "Severity scale used for this study, adapted to the organisation's context; each feared event is positioned on it (cf. Feared events)."))
    _table([T("Niveau", "Level"), T("Libellé", "Label"), T("Description", "Description"), T("Impact financier", "Financial impact"), T("Impact réputation", "Reputation impact"), T("Impact réglementaire", "Regulatory impact"), T("Impact données perso.", "Personal-data impact"), T("Impact opérationnel", "Operational impact")],
           "gravity_scale", ["niveau", "label", "description", "impact_financier", "impact_reputation", "impact_reglementaire", "impact_donnees_perso", "impact_operationnel"],
           T("Échelle de gravité et critères d'impact", "Severity scale and impact criteria"))
    portrait()
    hb = doc.add_heading(T("Annexe B — Niveaux de risque (gravité × vraisemblance)", "Appendix B — Risk levels (severity × likelihood)"), level=2); no_number(hb); bookmark(hb, "annexe_risque")
    I(T("Matrice d'acceptabilité du risque paramétrée dans l'étude : à chaque couple (gravité, vraisemblance) correspond un niveau de risque.",
        "Risk acceptability matrix configured in the study: each (severity, likelihood) pair maps to a risk level."))
    IMG("@@IMG:risk_matrix_ref@@", T("Matrice d'acceptabilité du risque (gravité × vraisemblance)", "Risk acceptability matrix (severity × likelihood)"))

    # ===== Native automatic heading numbering (multilevel list) =====
    ABSTRACT = (
        '<w:abstractNum %s w:abstractNumId="50"><w:multiLevelType w:val="multilevel"/>'
        '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%%1."/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="0" w:hanging="0"/></w:pPr></w:lvl>'
        '<w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%%1.%%2"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="0" w:hanging="0"/></w:pPr></w:lvl>'
        '<w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%%1.%%2.%%3"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="0" w:hanging="0"/></w:pPr></w:lvl>'
        '</w:abstractNum>' % nsdecls('w')
    )
    numbering = doc.part.numbering_part.element
    _nums = numbering.findall(qn('w:num'))
    _abs = parse_xml(ABSTRACT)
    if _nums: _nums[0].addprevious(_abs)
    else: numbering.append(_abs)
    numbering.append(parse_xml('<w:num %s w:numId="50"><w:abstractNumId w:val="50"/></w:num>' % nsdecls('w')))
    for _name, _ilvl in (("Heading 1", 0), ("Heading 2", 1), ("Heading 3", 2)):
        _pPr = doc.styles[_name].element.get_or_add_pPr()
        _pPr.append(parse_xml('<w:numPr %s><w:ilvl w:val="%d"/><w:numId w:val="50"/></w:numPr>' % (nsdecls('w'), _ilvl)))
    doc.settings.element.append(parse_xml('<w:updateFields %s w:val="true"/>' % nsdecls('w')))

    # ===== Footer on every section (suppressed on the cover) =====
    footer_title = T("Rapport d'analyse des risques — EBIOS Risk Manager", "Risk Analysis Report — EBIOS Risk Manager")
    for _sec in doc.sections:
        _set_footer(_sec, footer_title)
    doc.sections[0].different_first_page_header_footer = True
    doc.sections[0].first_page_footer.is_linked_to_previous = False

    out = os.path.join(TPL_DIR, "ebios-report-%s.docx" % LANG)
    doc.save(out)
    print("template écrit:", os.path.basename(out), round(os.path.getsize(out) / 1024, 1), "Ko")


for _lang in ("fr", "en"):
    build(_lang)
